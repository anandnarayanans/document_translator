import os
import threading
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import pdfplumber
import arabic_reshaper
from bidi.algorithm import get_display
from transformers import MarianMTModel, MarianTokenizer
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.utils import ImageReader
from io import BytesIO
import re
import tempfile
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert as docx2pdf_convert
from pdf2docx import Converter
# import pythoncom
from datetime import datetime
import json
import random
import string
from langdetect import detect, DetectorFactory
from langcodes import Language
import easyocr
from PIL import Image, ImageDraw, ImageFont
import numpy as np
from io import BytesIO

from pyngrok import ngrok


app = Flask(__name__)
CORS(app)
 # Set seed for consistent language detection
DetectorFactory.seed = 0

translations = {
    "1": {
        "file_name": "example.pdf",
        "translation_date": "2024-08-02",
        "translation_time": "12:00:00"
    },
    "2": {
        "file_name": "another_example.pdf",
        "translation_date": "2024-08-02",
        "translation_time": "12:30:00"
    }
}
 
def generate_random_username():
    """Generate a random username."""
    return ''.join(random.choices(string.ascii_lowercase + string.digits, k=8))
 
def extract_text_images_tables_from_pdf(pdf_path):
    """Extract text, images, and tables from a PDF file."""
    extracted_content = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_content = {"text": "", "translated_text": "", "images": [], "tables": []}
 
            text = page.extract_text(x_tolerance=2, y_tolerance=2)
            if text:
                bidi_text = get_display(arabic_reshaper.reshape(text))
                page_content["text"] = bidi_text
 
       # Extract images with their positions
            for image in page.images:
                try:
                    img_bbox = (image["x0"], image["top"], image["x1"], image["bottom"])
                    img_obj = page.within_bbox(img_bbox).to_image(resolution=300)
                    img_data = BytesIO()
                    img_obj.save(img_data, format='PNG')
                    img_data.seek(0)
                    img_bytes = img_data.read()
                    # Store both position and image data
                    page_content["images"].append({
                        "bbox": img_bbox,
                        "image": img_bytes,
                        "width": image["width"],
                        "height": image["height"]
                    })
                except Exception as e:
                    print(f"Error processing image: {e}")
                    continue
 
            tables = page.extract_tables()
            for table in tables:
                reshaped_table = []
                for row in table:
                    reshaped_row = [get_display(arabic_reshaper.reshape(cell)) if cell else '' for cell in row]
                    reshaped_table.append(reshaped_row)
                page_content["tables"].append(reshaped_table)
 
            extracted_content.append(page_content)
    return extracted_content
 
def translate_arabic_to_english(model, tokenizer, lines):
    """Translate Arabic text to English using MarianMT model."""
    translated_lines = []
    for line in lines:
        translated = model.generate(**tokenizer(line, return_tensors="pt", padding=True))
        translated_text = tokenizer.batch_decode(translated, skip_special_tokens=True)[0]
        translated_lines.append(translated_text)
    return translated_lines
 
def get_file_size(file_path):
    """ Get the file size in MB. """
    file_size_mb = os.path.getsize(file_path) / (1024 * 1024)  # Convert to MB
    return round(file_size_mb, 2)

def get_number_of_pages(file_path):
    """ Get the number of pages using pdfplumber. """
    with pdfplumber.open(file_path) as pdf:
        return len(pdf.pages)
    
def get_project(file_name):
    if "Guidelines" in file_name:
        return "Policies/Guidelines"
    else:
        return "Proposals"
    
def get_full_language_name(detected_lang):
    try:
        return Language.get(detected_lang).display_name().capitalize()
    except:
        return detected_lang.capitalize()
    
def clean_translation(text):
    """Clean the translated text to remove unwanted artifacts."""
    text = text.replace("Ta &apos; air", "Name")
    text = text.replace("Emil", "Email")
    text = text.replace("date", "Date")
    text = text.replace("Page 1 of 2", "")
    text = text.replace("Page 2 of 2", "")
    text = text.replace("Page 2 of 2", "").replace("land,","").replace("&apos; s", "").replace("hey,","").replace("Hey,","").replace("H-E-E-E-E-E-E-E-E","").replace("Arr-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L","").replace("land,","").replace("Argh!","").replace("Oh, my God.","775").replace("&gt","").replace("oh,","").replace("%20","20%").replace("...","").replace("L-L-L-L-L-L-L-L","").replace(",,,","")
    text = text.replace("&apos;","").replace("-------L-L-L-L-L","").replace("Temporarily hysterical","Food allowance").replace("Screw it","Medical allowance").replace("oh.","").replace(", , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , ,","push four")
    text = text.replace("Y-","").replace("R","").replace("Y","")
    text = text.replace("LA","Date")
   
    cleaned_text = re.sub(r'\b(?:no|not|non-|, ,)\b', '', text, flags=re.IGNORECASE).strip()
    return cleaned_text

def trans_image_ocr(image_bytes,model,tokenizer):
    """Perform OCR on an image, handle both Arabic and English text with proper formatting."""
    try:
        reader = easyocr.Reader(['ar', 'en'])
        img_stream = BytesIO(image_bytes)
        image = Image.open(img_stream).convert("RGB")
        
        # Perform OCR
        results = reader.readtext(np.array(image))
        
        # Create a drawing object
        draw = ImageDraw.Draw(image)
        
        # Common substitutions for OCR mistakes
        ocr_corrections = {
            'R5YAL': 'ROYAL',
            'C5MMISSI5N': 'COMMISSION',
            'M4KK4H': 'MAKKAH',
            'CITV': 'CITY',
            '4NL': 'AND',
            'HLY': 'HOLY',
            '5': 'O',
            '4': 'A',
            'V': 'Y'
        }
        
        all_translated_text = []
        contains_text = False
        
        for bbox, text, confidence in results:
            contains_text = True
            
            # Check if text contains Arabic characters
            if any('\u0600' <= char <= '\u06FF' for char in text):
                # Translate Arabic text to English
                final_text = translate_arabic_to_english(model, tokenizer, [text])[0]
            else:
                # Apply global replacements for known OCR issues
                for wrong, right in ocr_corrections.items():
                    text = text.replace(wrong, right)  # Replace full patterns globally
                
                # Replace any remaining '5' with 'O' in the entire string
                text = text.replace('5', 'O')

                # Ensure final corrections are applied at the character level
                corrected_text = ''
                for char in text:
                    corrected_text += ocr_corrections.get(char, char)
                
                final_text = corrected_text
                
                # Special case for known phrases
                if any(phrase in final_text for phrase in ['ROYAL COMMISSION', 'MAKKAH CITY', 'HOLY SITES']):
                    # Preserve exact casing and formatting for official names
                    final_text = final_text.upper()
                else:
                    # For other English text, ensure proper capitalization
                    final_text = ' '.join(word.capitalize() for word in final_text.split())
            
            all_translated_text.append(final_text)
            
            # Get coordinates for text replacement
            x_min = min(p[0] for p in bbox)
            y_min = min(p[1] for p in bbox)
            x_max = max(p[0] for p in bbox)
            y_max = max(p[1] for p in bbox)
            
            # Create white background for clarity
            draw.rectangle([x_min, y_min, x_max, y_max], fill="white")
            
            # Calculate original text height to maintain font size
            text_height = y_max - y_min
            font_size = int(text_height * 0.62)  # Adjust this factor if needed
            
            # Draw text with consistent font
            try:
                font = ImageFont.truetype("Helvetica.ttf", font_size)
            except:
                font = ImageFont.load_default()
            
            # Draw the text
            draw.text((x_min, y_min), final_text, fill="black", font=font)
        
        # Return original image if no text was found
        if not contains_text:
            return image_bytes, ""
        
        # Convert modified image back to bytes
        output = BytesIO()
        image.save(output, format='PNG')
        output.seek(0)
        
        return output.read(), ""
    
    except Exception as e:
        print(f"Error in trans_image_ocr: {e}")
        return image_bytes, ""  #

def save_text_images_tables_to_pdf(pages_content, output_pdf_path, model, tokenizer):
    """Save the given text, images, and tables to a PDF file."""
    c = canvas.Canvas(output_pdf_path, pagesize=letter)
    width, height = letter
    styles = getSampleStyleSheet()
    styleN = styles['Normal']
    left_margin, right_margin = 35, 35
    current_font_size = 11
    
    for page_content in pages_content:
        y = height - 40
        
        # Process and draw images with translated text
        for img_data in page_content["images"]:
            # Get the original image data and position
            img_bytes = img_data["image"]
            img_bbox = img_data["bbox"]
            
            # Translate and modify the image
            modified_img_bytes, _  = trans_image_ocr(img_bytes,model,tokenizer)
            
            # Draw the modified image
            try:

                                # Create PIL Image from modified bytes
                img = Image.open(BytesIO(modified_img_bytes))
                

                # Convert PIL Image to format compatible with ReportLab
                img_stream = BytesIO()
                img.save(img_stream, format='PNG')
                img_stream.seek(0)
                img = ImageReader(img_stream)
    
                 # Calculate positions maintaining aspect ratio
                x0, y0, x1, y1 = img_bbox
                img_width = x1 - x0
                img_height = y1 - y0
                
                # Convert to PDF coordinates (bottom-left origin)
                y_draw = height - y1  # Adjust for PDF coordinate system
                
                # Draw image with proper positioning and scaling
                c.drawImage(img, 
                          x0,                    # x position 
                          y_draw,                # y position
                          width=img_width,       # maintain original width
                          height=img_height,     # maintain original height
                          mask='auto'            # handle transparency
                )
                y -= img_height + 15
                
            except Exception as e:
                print(f"Error drawing image: {e}")
                continue

        c.setFont("Helvetica", current_font_size)
        translated_text = page_content["translated_text"]
        translated_lines = translated_text.split('\n')
 
        # Variables for list detection
        in_numbered_list = False
        in_address_list = False
        next_line_starts_address_list = False

        i = 0
        while i < len(translated_lines):
            line = translated_lines[i]
            
            # Check for numbered list start
            starts_with_number = line.strip() and line.strip()[0].isdigit() and '.' in line.strip()[:3]

            # Check for address-based list start in next line
            if "address" in line.lower() and ":" in line:
                if i > 0:
                    y -= 15  # Line break above the sentence before the address-based list
                next_line_starts_address_list = True
            elif next_line_starts_address_list:
                in_address_list = True
                next_line_starts_address_list = False
            
            # Detect sentences with ':' and '-' and handle line breaks
            if ":" in line and "-" in line:
                if i < len(translated_lines) - 1 and translated_lines[i + 1].strip()[0].isdigit():
                    y -= 15  # Add a line break above the sentence

            # Handle list end conditions
            if in_numbered_list and not starts_with_number:
                in_numbered_list = False
                y -= 15  # Line break above detected list
            
            if in_address_list:
                words = line.split()
                if len(words) >= 14:  # Check for long sentence
                    in_address_list = False
                    y -= 15  # Line break above detected list
            
            # Check if the line contains a colon and handle it as a full sentence if so
            if ":" in line:
                colon_idx = line.index(":")
                sentence = line[:colon_idx + 1].strip().capitalize()
                remaining_text = line[colon_idx + 1:].strip()
               
                # Draw the bold sentence with word wrapping and a single space after the colon
                c.setFont("Helvetica-Bold", current_font_size)
                sentence_to_print = (sentence + ":" if sentence[-1] != ":" else sentence) + " "
               
                # Word wrapping for the bold sentence
                text_offset = left_margin
                for word in sentence_to_print.split():
                    if text_offset + c.stringWidth(word + " ") > (width - right_margin):
                        y -= 15
                        text_offset = left_margin
                    c.drawString(text_offset, y, word)
                    text_offset += c.stringWidth(word + " ")
 
                # Draw remaining text in normal font with word wrapping
                c.setFont("Helvetica", current_font_size)
                for word in remaining_text.split():
                    if text_offset + c.stringWidth(word + " ") > (width - right_margin):
                        y -= 15
                        text_offset = left_margin
                    c.drawString(text_offset, y, word)
                    text_offset += c.stringWidth(word + " ")
            else:
                # Draw the line as normal with word wrapping within the margins
                text_offset = left_margin
                for word in line.split():
                    if text_offset + c.stringWidth(word + " ") > (width - right_margin):
                        y -= 15
                        text_offset = left_margin
                    c.drawString(text_offset, y, word)
                    text_offset += c.stringWidth(word + " ")
 
            y -= 15  # Minimal space to the next line of text
 
            # Check for page break and prevent empty page creation
            if y < 40:
                if y == height - 40:  # Ensure the current page is not empty
                    break
                c.showPage()
                y = height - 40
                c.setFont("Helvetica", current_font_size)  # Maintain font size across pages
            
            i += 1
 
        # Draw tables with word wrapping in each cell
        for table in page_content["tables"]:
            table_style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.white),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('TOPPADDING', (0, 0), (-1, -1), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('WORDWRAP', (0, 0), (-1, -1), 'CJK')
            ])
 
            wrapped_table = []
            for row in table:
                wrapped_row = [Paragraph(cell, styleN) for cell in row]
                wrapped_table.append(wrapped_row)
 
            col_widths = [(width - left_margin - right_margin) / len(wrapped_table[0])] * len(wrapped_table[0])
            t = Table(wrapped_table, colWidths=col_widths)
            t.setStyle(table_style)
            w, h = t.wrap(width, y)
            if h > y:
                if y != height - 40:  # Avoid adding a new page if the current one is empty
                    c.showPage()
                y = height - 40
                c.setFont("Helvetica", current_font_size)  # Maintain font size across pages
            t.drawOn(c, left_margin, y - h)
            y -= h + 15
 
        if y != height - 40:  # Avoid adding unnecessary blank pages
            c.showPage()
 
    c.save()


def save_text_images_tables_to_docx(pages_content, output_docx_path):
    """Save the given text, images, and tables to a DOCX file."""
    doc = Document()
 
    for page_content in pages_content:
        translated_text = page_content["translated_text"]
        translated_lines = translated_text.split('\n')
 
        for line in translated_lines:
            p = doc.add_paragraph(line)
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
 
        for table in page_content["tables"]:
            t = doc.add_table(rows=len(table), cols=len(table[0]))
            t.style = 'Table Grid'
            for row_idx, row in enumerate(table):
                for col_idx, cell_text in enumerate(row):
                    cell = t.cell(row_idx, col_idx)
                    cell.text = cell_text
 
    doc.save(output_docx_path)
 
 
 
def translate_pdf(file_path, translation_id, initial_format):
    
     # First, extract only the first page for language detection
    with pdfplumber.open(file_path) as pdf:
        first_page = pdf.pages[0]
        first_page_text = first_page.extract_text()
        if not first_page_text:
            raise ValueError("No text found in the first page")
        
        # Detect language only from first page
        detected_lang = detect(first_page_text)
        model_name = f"Helsinki-NLP/opus-mt-{detected_lang}-en"
        tokenizer = MarianTokenizer.from_pretrained(model_name)
        model = MarianMTModel.from_pretrained(model_name)
        
    extracted_content = extract_text_images_tables_from_pdf(file_path)
  
    for page_content in extracted_content:
        lines = page_content["text"].split('\n')
        translated_lines = translate_arabic_to_english(model, tokenizer, lines)
        cleaned_lines = [clean_translation(line) for line in translated_lines]
        page_content["translated_text"] = "\n".join(cleaned_lines)
 
        translated_tables = []
        for table in page_content["tables"]:
            translated_table = []
            for row in table:
                translated_row = translate_arabic_to_english(model, tokenizer, row)
                cleaned_row = [clean_translation(cell) for cell in translated_row]
                translated_table.append(cleaned_row)
            translated_tables.append(translated_table)
        page_content["tables"] = translated_tables
 
    output_dir = tempfile.mkdtemp()
    output_pdf_path = os.path.join(output_dir, "translated_output.pdf")
    save_text_images_tables_to_pdf(extracted_content, output_pdf_path, model, tokenizer)
 
    translation_details = {
        "status": "completed",
        "file_path": output_pdf_path,
        "download_link": output_pdf_path,
        "initial_format": initial_format,
        "file_name": os.path.basename(file_path),
        "translation_time": datetime.now().strftime("%H:%M:%S"),
        "translation_date": datetime.now().strftime("%Y-%m-%d"),
         "number_of_pages": get_number_of_pages(file_path),
        "file_size": get_file_size(file_path),
        # "language": "Arabic-English",
        "language": f"{get_full_language_name(detected_lang)}-English",
        "project": get_project(os.path.basename(file_path))
    }
    translations[translation_id] = translation_details
    save_translations(translations)
 
 

@app.route('/upload', methods=['POST'])
def upload():
    # pythoncom.CoInitialize()
    """Handle file upload and initiate the translation process."""
    skip_warning = request.headers.get('ngrok-skip-browser-warning')

    file = request.files['file']
    if file:
        temp_dir = tempfile.mkdtemp()
        file_path = os.path.join(temp_dir, file.filename)
        file.save(file_path)
 
        initial_format = 'pdf' if file.filename.endswith('.pdf') else 'docx'

        return jsonify({'message': 'File uploaded successfully', 'file_path': file_path, 'initial_format': initial_format})
    return jsonify({'message': 'No file uploaded'}), 400
 
@app.route('/translate', methods=['POST'])
def translate():
    skip_warning = request.headers.get('ngrok-skip-browser-warning')
    """Start the translation process."""
    data = request.get_json()
    file_path = data.get('file_path')
    initial_format = data.get('initial_format')
 
    if file_path:
        translation_id = str(len(translations) + 1)
        translations[translation_id] = {"status": "in_progress", "file_path": None}
 
        translation_thread = threading.Thread(target=translate_pdf, args=(file_path, translation_id, initial_format))
        translation_thread.start()
 
        return jsonify({'message': 'Translation started', 'translation_id': translation_id})
    return jsonify({'message': 'File path not provided'}), 400
 
@app.route('/translation_status/<translation_id>', methods=['GET'])
def translation_status(translation_id):
    skip_warning = request.headers.get('ngrok-skip-browser-warning')
    """Check the status of the translation."""
    translation = translations.get(translation_id)
    if translation:
        return jsonify({'status': translation['status'], 'file_path': translation['file_path']})
    return jsonify({'status': 'not_found'}), 404
 
@app.route('/download/<translation_id>', methods=['GET'])
def download_translated(translation_id):
    skip_warning = request.headers.get('ngrok-skip-browser-warning')
    """Download the translated document in the initial format."""
    translation = translations.get(translation_id)
    if translation and translation['status'] == 'completed':
        file_path = translation['file_path']
        initial_format = translation.get('initial_format')
 

        return send_file(file_path, as_attachment=True)
    return jsonify({'message': 'File not found or translation not completed'}), 404
 
@app.route('/preview/<translation_id>', methods=['GET'])
def download_pdf(translation_id):
    """Download the translated PDF file."""
    skip_warning = request.headers.get('ngrok-skip-browser-warning')
    translation = translations.get(translation_id)
    if translation and translation['status'] == 'completed':
        file_path = translation['file_path']
        return send_file(file_path, as_attachment=False)
    return jsonify({'message': 'File not found or translation not completed'}), 404
 
@app.route('/translations', methods=['GET'],strict_slashes=False)
def get_translations():
    """Get the list of translations."""
    skip_warning = request.headers.get('ngrok-skip-browser-warning')
    return jsonify(list(translations.values()))
 
TRANSLATIONS_FILE = 'translations.json'
 
def load_translations():
    """Load translations from the JSON file."""
    if os.path.exists(TRANSLATIONS_FILE):
        with open(TRANSLATIONS_FILE, 'r') as file:
            return json.load(file)
    return {}
 
def save_translations(translations):
    """Save translations to the JSON file."""
    with open(TRANSLATIONS_FILE, 'w') as file:
        json.dump(translations, file)
 
@app.route('/download/<translation_id>', methods=['GET'])
def download_file(translation_id):
    skip_warning = request.headers.get('ngrok-skip-browser-warning')
    translation = translations.get(translation_id)
    if translation and 'file_path' in translation:
        file_path = translation['file_path']
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True)
    return jsonify({"message": "File not found or translation not completed"}), 404
 
translations = load_translations()
 
# Run ngrok for public access
public_url = ngrok.connect(1001)
print(f"Public URL: {public_url}")

# Run Flask server
threading.Thread(target=app.run, kwargs={"port": 1001}).start()