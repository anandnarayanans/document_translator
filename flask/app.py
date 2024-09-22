import os
import threading
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import pdfplumber
import arabic_reshaper
from bidi.algorithm import get_display
from transformers import MarianMTModel, MarianTokenizer
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.utils import ImageReader
from io import BytesIO
import re
import tempfile
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert as docx2pdf_convert
from pdf2docx import Converter
import pythoncom
from datetime import datetime
import json
import random
import string
 
app = Flask(__name__)
CORS(app)
 
translations = {
    "1": {
        "file_name": "example.pdf",
        "translation_date": "2024-08-02",
        "translation_time": "12:00:00"
    },
    "2": {
        "file_name": "another_example.pdf",
        "translation_date": "2024-08-02",
        "translation_time": "12:30:00"
    }
}
 
def generate_random_username():
    """Generate a random username."""
    return ''.join(random.choices(string.ascii_lowercase + string.digits, k=8))
 
def extract_text_images_tables_from_pdf(pdf_path):
    """Extract text, images, and tables from a PDF file."""
    extracted_content = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_content = {"text": "", "translated_text": "", "images": [], "tables": []}
 
            text = page.extract_text(x_tolerance=2, y_tolerance=2)
            if text:
                bidi_text = get_display(arabic_reshaper.reshape(text))
                page_content["text"] = bidi_text
 
            for image in page.images:
                try:
                    img_bbox = (image["x0"], image["top"], image["x1"], image["bottom"])
                    img_obj = page.within_bbox(img_bbox).to_image(resolution=300)
                    img_data = BytesIO()
                    img_obj.save(img_data, format='PNG')
                    img_data.seek(0)
                    img_bytes = img_data.read()
                    page_content["images"].append((img_bbox, img_bytes))
                except Exception as e:
                    print(f"Error processing image: {e}")
                    continue
 
            tables = page.extract_tables()
            for table in tables:
                reshaped_table = []
                for row in table:
                    reshaped_row = [get_display(arabic_reshaper.reshape(cell)) if cell else '' for cell in row]
                    reshaped_table.append(reshaped_row)
                page_content["tables"].append(reshaped_table)
 
            extracted_content.append(page_content)
    return extracted_content
 
def translate_arabic_to_english(model, tokenizer, lines):
    """Translate Arabic text to English using MarianMT model."""
    translated_lines = []
    for line in lines:
        translated = model.generate(**tokenizer(line, return_tensors="pt", padding=True))
        translated_text = tokenizer.batch_decode(translated, skip_special_tokens=True)[0]
        translated_lines.append(translated_text)
    return translated_lines
 
def get_file_size(file_path):
    """ Get the file size in MB. """
    file_size_mb = os.path.getsize(file_path) / (1024 * 1024)  # Convert to MB
    return round(file_size_mb, 2)

def get_number_of_pages(file_path):
    """ Get the number of pages using pdfplumber. """
    with pdfplumber.open(file_path) as pdf:
        return len(pdf.pages)
    
def clean_translation(text):
    """Clean the translated text to remove unwanted artifacts."""
    text = text.replace("Ta &apos; air", "Name")
    text = text.replace("Emil", "Email")
    text = text.replace("date", "Date")
    text = text.replace("Page 1 of 2", "")
    text = text.replace("Page 2 of 2", "")
    text = text.replace("Page 2 of 2", "").replace("land,","").replace("&apos; s", "").replace("hey,","").replace("Hey,","").replace("H-E-E-E-E-E-E-E-E","").replace("Arr-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L-L","").replace("land,","").replace("Argh!","").replace("Oh, my God.","775").replace("&gt","").replace("oh,","").replace("%20","20%").replace("...","").replace("L-L-L-L-L-L-L-L","").replace(",,,","")
    text = text.replace("&apos;","").replace("-------L-L-L-L-L","").replace("Temporarily hysterical","Food allowance").replace("Screw it","Medical allowance").replace("oh.","").replace(", , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , , ,","push four")
    text = text.replace("Y-","").replace("R","").replace("Y","")
    text = text.replace("LA","Date")
   
    cleaned_text = re.sub(r'\b(?:no|not|non-|, ,)\b', '', text, flags=re.IGNORECASE).strip()
    return cleaned_text
 
def save_text_images_tables_to_pdf(pages_content, output_pdf_path, model, tokenizer):
    """Save the given text, images, and tables to a PDF file."""
    c = canvas.Canvas(output_pdf_path, pagesize=letter)
    width, height = letter
    styles = getSampleStyleSheet()
    styleN = styles['Normal']
 
    for page_content in pages_content:
        y = height - 40  # Start from top of the page
 
        # Draw images first to avoid overlapping text
        for img_bbox, img_bytes in page_content["images"]:
            x0, top, x1, bottom = img_bbox
            image_stream = BytesIO(img_bytes)
            img_height = bottom - top
            img_width = x1 - x0
            c.drawImage(ImageReader(image_stream), x0, height - bottom, width=img_width, height=img_height)
            y -= img_height + 15  # Add some extra space between image and text
 
        c.setFont("Helvetica", 11)
 
        translated_text = page_content["translated_text"]
        translated_lines = translated_text.split('\n')
 
        for line in translated_lines:
            words = line.split()
            current_line = ""
            for word in words:
                if c.stringWidth(current_line + " " + word) > (width - 80):  # 80 is a margin value
                    c.drawString(40, y, current_line.strip())
                    y -= 15  # Move to the next line with a fixed spacing
                    current_line = ""
                if current_line == "":
                    current_line = word
                else:
                    current_line += " " + word
            if current_line:
                c.drawString(40, y, current_line.strip())
                y -= 15  # Move to the next line with a fixed spacing
 
            if y < 40:  # Create a new page if we reach the bottom of the current page
                c.showPage()
                y = height - 40
 
        # Draw tables
        for table in page_content["tables"]:
            table_style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.white),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('TOPPADDING', (0, 0), (-1, -1), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('WORDWRAP', (0, 0), (-1, -1), 'CJK')
            ])
 
            wrapped_table = []
            for row in table:
                wrapped_row = [Paragraph(cell, styleN) for cell in row]
                wrapped_table.append(wrapped_row)
 
            col_widths = [(width - 80) / len(wrapped_table[0])] * len(wrapped_table[0])  # Adjust column width to fit page
            t = Table(wrapped_table, colWidths=col_widths)
            t.setStyle(table_style)
            w, h = t.wrap(width, y)
            if h > y:
                c.showPage()
                y = height - 40
            t.drawOn(c, 40, y - h)
            y -= h + 15
 
        c.showPage()  # Create a new page for the next set of content
 
    c.save()
 
def save_text_images_tables_to_docx(pages_content, output_docx_path):
    """Save the given text, images, and tables to a DOCX file."""
    doc = Document()
 
    for page_content in pages_content:
        translated_text = page_content["translated_text"]
        translated_lines = translated_text.split('\n')
 
        for line in translated_lines:
            p = doc.add_paragraph(line)
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
 
        for table in page_content["tables"]:
            t = doc.add_table(rows=len(table), cols=len(table[0]))
            t.style = 'Table Grid'
            for row_idx, row in enumerate(table):
                for col_idx, cell_text in enumerate(row):
                    cell = t.cell(row_idx, col_idx)
                    cell.text = cell_text
 
    doc.save(output_docx_path)
 
 
 
def translate_pdf(file_path, translation_id, initial_format):
    extracted_content = extract_text_images_tables_from_pdf(file_path)
    model_name = "Helsinki-NLP/opus-mt-ar-en"
    tokenizer = MarianTokenizer.from_pretrained(model_name)
    model = MarianMTModel.from_pretrained(model_name)
 
    for page_content in extracted_content:
        lines = page_content["text"].split('\n')
        translated_lines = translate_arabic_to_english(model, tokenizer, lines)
        cleaned_lines = [clean_translation(line) for line in translated_lines]
        page_content["translated_text"] = "\n".join(cleaned_lines)
 
        translated_tables = []
        for table in page_content["tables"]:
            translated_table = []
            for row in table:
                translated_row = translate_arabic_to_english(model, tokenizer, row)
                cleaned_row = [clean_translation(cell) for cell in translated_row]
                translated_table.append(cleaned_row)
            translated_tables.append(translated_table)
        page_content["tables"] = translated_tables
 
    output_dir = tempfile.mkdtemp()
    output_pdf_path = os.path.join(output_dir, "translated_output.pdf")
    save_text_images_tables_to_pdf(extracted_content, output_pdf_path, model, tokenizer)
 
    translation_details = {
        "status": "completed",
        "file_path": output_pdf_path,
        "download_link": output_pdf_path,
        "initial_format": initial_format,
        "file_name": os.path.basename(file_path),
        "translation_time": datetime.now().strftime("%H:%M:%S"),
        "translation_date": datetime.now().strftime("%Y-%m-%d"),
         "number_of_pages": get_number_of_pages(file_path),
        "file_size": get_file_size(file_path),
        "language": "Arabic-English"
    }
    translations[translation_id] = translation_details
    save_translations(translations)
 
 
 
 
def convert_docx_to_pdf(docx_path):
    """Convert DOCX to PDF using the docx2pdf library."""
    output_pdf_path = docx_path.replace('.docx', '.pdf')
    docx2pdf_convert(docx_path, output_pdf_path)
    return output_pdf_path
 
def convert_pdf_to_docx(pdf_path):
    """Convert PDF to DOCX using the pdf2docx library."""
    output_docx_path = pdf_path.replace('.pdf', '.docx')
    cv = Converter(pdf_path)
    cv.convert(output_docx_path, start=0, end=None)
    cv.close()
    return output_docx_path
 
@app.route('/upload', methods=['POST'])
def upload():
    pythoncom.CoInitialize()
    """Handle file upload and initiate the translation process."""
    file = request.files['file']
    if file:
        temp_dir = tempfile.mkdtemp()
        file_path = os.path.join(temp_dir, file.filename)
        file.save(file_path)
 
        initial_format = 'pdf' if file.filename.endswith('.pdf') else 'docx'
       
        if initial_format == 'docx':
            file_path = convert_docx_to_pdf(file_path)
 
        return jsonify({'message': 'File uploaded successfully', 'file_path': file_path, 'initial_format': initial_format})
    return jsonify({'message': 'No file uploaded'}), 400
 
@app.route('/translate', methods=['POST'])
def translate():
    """Start the translation process."""
    data = request.get_json()
    file_path = data.get('file_path')
    initial_format = data.get('initial_format')
 
    if file_path:
        translation_id = str(len(translations) + 1)
        translations[translation_id] = {"status": "in_progress", "file_path": None}
 
        translation_thread = threading.Thread(target=translate_pdf, args=(file_path, translation_id, initial_format))
        translation_thread.start()
 
        return jsonify({'message': 'Translation started', 'translation_id': translation_id})
    return jsonify({'message': 'File path not provided'}), 400
 
@app.route('/translation_status/<translation_id>', methods=['GET'])
def translation_status(translation_id):
    """Check the status of the translation."""
    translation = translations.get(translation_id)
    if translation:
        return jsonify({'status': translation['status'], 'file_path': translation['file_path']})
    return jsonify({'status': 'not_found'}), 404
 
@app.route('/download/<translation_id>', methods=['GET'])
def download_translated(translation_id):
    """Download the translated document in the initial format."""
    translation = translations.get(translation_id)
    if translation and translation['status'] == 'completed':
        file_path = translation['file_path']
        initial_format = translation.get('initial_format')
 
        if initial_format == 'docx' and file_path.endswith('.pdf'):
            file_path = convert_pdf_to_docx(file_path)
        elif initial_format == 'pdf' and file_path.endswith('.docx'):
            file_path = convert_docx_to_pdf(file_path)
           
        return send_file(file_path, as_attachment=True)
    return jsonify({'message': 'File not found or translation not completed'}), 404
 
@app.route('/preview/<translation_id>', methods=['GET'])
def download_pdf(translation_id):
    """Download the translated PDF file."""
    translation = translations.get(translation_id)
    if translation and translation['status'] == 'completed':
        file_path = translation['file_path']
        return send_file(file_path, as_attachment=False)
    return jsonify({'message': 'File not found or translation not completed'}), 404
 
@app.route('/translations', methods=['GET'])
def get_translations():
    """Get the list of translations."""
    return jsonify(list(translations.values()))
 
TRANSLATIONS_FILE = 'translations.json'
 
def load_translations():
    """Load translations from the JSON file."""
    if os.path.exists(TRANSLATIONS_FILE):
        with open(TRANSLATIONS_FILE, 'r') as file:
            return json.load(file)
    return {}
 
def save_translations(translations):
    """Save translations to the JSON file."""
    with open(TRANSLATIONS_FILE, 'w') as file:
        json.dump(translations, file)
 
@app.route('/download/<translation_id>', methods=['GET'])
def download_file(translation_id):
    translation = translations.get(translation_id)
    if translation and 'file_path' in translation:
        file_path = translation['file_path']
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True)
    return jsonify({"message": "File not found or translation not completed"}), 404
 
translations = load_translations()
 
 
if __name__ == '__main__':
    app.run(debug=False, use_reloader=False, port=5001)